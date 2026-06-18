"""
Resume Tailor — rewrites resume bullet points per job description and generates PDF.

Pipeline:
  1. Takes CandidateProfile + MatchResult (has skill_overlap, skill_gaps, reasoning)
  2. For each work experience entry, LLM rewrites 3-5 bullet points
     emphasising overlapping skills from the job
  3. Generates a PDF via python-docx → WeasyPrint (same pipeline as cover letters)
  4. Returns path to the tailored resume PDF
"""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Optional

from backend.config import settings
from backend.models.profile import CandidateProfile, MatchResult

logger = logging.getLogger(__name__)

TAILORED_DIR = settings.data_dir / "tailored_resumes"

BULLET_REWRITE_PROMPT = """\
You are a professional resume writer. Rewrite the following bullet points from a
candidate's resume to better match a specific job description.

CANDIDATE
=========
Name: {name}
Role at this company: {title} @ {company}
Years at role: {years}

ORIGINAL BULLET POINTS:
{bullets}

JOB DESCRIPTION
===============
Company: {job_company}
Title: {job_title}
Description: {job_description}

KEY SKILLS THE JOB REQUIRES (emphasise these):
{overlap_skills}

SKILLS THE CANDIDATE SHOULD HIGHLIGHT IF POSSIBLE:
{gap_skills}

INSTRUCTIONS
============
Rewrite each bullet point to emphasise the key skills the job requires.
Keep it factual — do not invent experience.
Make each bullet specific, quantified where possible, and action-oriented.
Use strong action verbs.
Return ONLY the rewritten bullet points, one per line, starting with "• ".
Return 3-5 bullet points.

REWRITTEN BULLET POINTS:
"""


def _call_llm(prompt: str, system_prompt: str = "") -> str | None:
    """Call Claude API (primary) or OpenAI (fallback)."""
    if settings.claude_api_key:
        try:
            import httpx
            resp = httpx.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": settings.claude_api_key,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
                json={
                    "model": "claude-sonnet-4-20250514",
                    "max_tokens": 1500,
                    "system": system_prompt or "You are a professional resume writer.",
                    "messages": [{"role": "user", "content": prompt}],
                },
                timeout=60,
            )
            resp.raise_for_status()
            data = resp.json()
            return data["content"][0]["text"]
        except Exception as e:
            logger.warning("Claude API call failed: %s", e)

    if settings.openai_api_key:
        try:
            import httpx
            resp = httpx.post(
                "https://api.openai.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {settings.openai_api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "gpt-4o",
                    "max_tokens": 1500,
                    "messages": [
                        {"role": "system", "content": system_prompt or "You are a professional resume writer."},
                        {"role": "user", "content": prompt},
                    ],
                },
                timeout=60,
            )
            resp.raise_for_status()
            data = resp.json()
            return data["choices"][0]["message"]["content"]
        except Exception as e:
            logger.warning("OpenAI API call failed: %s", e)

    logger.warning("No LLM API key configured — resume tailoring unavailable")
    return None


def _generate_resume_docx(
    profile: CandidateProfile,
    rewritten_entries: dict[str, list[str]],
    output_path: Path,
) -> Path:
    """Generate a .docx resume with rewritten bullet points, then convert to PDF."""
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()

    # Name & contact
    title_style = doc.styles["Title"]
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    doc.add_paragraph(profile.full_name or "Candidate", style="Title")

    contact_line = ", ".join(
        p for p in [profile.email, profile.phone, profile.location, profile.linkedin_url, profile.github_url]
        if p
    )
    if contact_line:
        p = doc.add_paragraph(contact_line)
        p.style.font.size = Pt(10)

    # Summary
    s = profile.summary()
    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(
        f"{s.get('title', 'Professional')} with {profile.years_of_experience:.0f} years "
        f"of experience. Skilled in {', '.join(profile.skill_names[:8])}."
    )

    # Experience
    doc.add_heading("Experience", level=2)
    for exp in profile.experiences:
        doc.add_heading(f"{exp.title} at {exp.company}", level=3)
        if exp.start_date or exp.end_date:
            doc.add_paragraph(f"{exp.start_date or ''} — {exp.end_date or 'Present'}")

        # Use rewritten bullets if available
        entry_key = f"{exp.company}|{exp.title}"
        bullets = rewritten_entries.get(entry_key)
        if bullets:
            for b in bullets:
                doc.add_paragraph(b, style="List Bullet")
        elif exp.description:
            for line in exp.description.split("\n"):
                line = line.strip().lstrip("•-*")
                if line:
                    doc.add_paragraph(f"• {line}", style="List Bullet")

    # Skills
    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(profile.skill_names[:15]))

    # Education
    if profile.education:
        doc.add_heading("Education", level=2)
        for edu in profile.education:
            parts = [p for p in [edu.degree, edu.field, edu.institution] if p]
            doc.add_paragraph(" — ".join(parts))

    doc.save(str(output_path))
    logger.info("Tailored resume .docx saved to %s", output_path)
    return output_path


def tailor_resume(
    profile: CandidateProfile,
    match: MatchResult,
) -> Path | None:
    """Tailor resume bullet points for a specific job and return the PDF path."""
    TAILORED_DIR.mkdir(parents=True, exist_ok=True)

    if not profile.experiences:
        logger.warning("No experience entries to tailor")
        return None

    rewritten_entries: dict[str, list[str]] = {}
    job = match.job
    overlap_str = ", ".join(match.skill_overlap[:10]) if match.skill_overlap else "N/A"
    gap_str = ", ".join(match.skill_gaps[:10]) if match.skill_gaps else "None"

    for exp in profile.experiences:
        bullets_text = exp.description.strip() or "N/A"
        # Split into individual bullet points
        bullet_lines = [
            b.strip().lstrip("•-* ").strip()
            for b in bullets_text.split("\n")
            if b.strip()
        ]
        if not bullet_lines:
            continue

        prompt = BULLET_REWRITE_PROMPT.format(
            name=profile.full_name or "Candidate",
            title=exp.title,
            company=exp.company,
            years=exp.years_at_role or "N/A",
            bullets="\n".join(f"- {b}" for b in bullet_lines),
            job_company=job.company,
            job_title=job.title,
            job_description=job.description[:2000] if job.description else "N/A",
            overlap_skills=overlap_str,
            gap_skills=gap_str,
        )

        response = _call_llm(prompt)
        if response:
            rewritten = [
                line.strip().lstrip("•-* ").strip()
                for line in response.strip().split("\n")
                if line.strip() and not line.strip().startswith(("```", "Here", "**"))
            ]
            entry_key = f"{exp.company}|{exp.title}"
            rewritten_entries[entry_key] = rewritten[:5]

    if not rewritten_entries:
        logger.warning("No bullet points were rewritten — using original CV")
        return None

    # Generate the document
    output_stem = f"tailored_{job.company.replace(' ', '_')}_{job.id[:8]}"
    docx_path = TAILORED_DIR / f"{output_stem}.docx"
    pdf_path = TAILORED_DIR / f"{output_stem}.pdf"

    _generate_resume_docx(profile, rewritten_entries, docx_path)

    # Convert to PDF via WeasyPrint
    try:
        from weasyprint import HTML
        from docx import Document as DocxReader

        docx_text = []
        doc = DocxReader(str(docx_path))
        for para in doc.paragraphs:
            docx_text.append(para.text)

        html_content = f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<style>
body {{ font-family: 'Space Grotesk', sans-serif; margin: 1.5in; font-size: 11pt; line-height: 1.5; }}
h1 {{ font-size: 20pt; margin-bottom: 0.25in; }}
h2 {{ font-size: 14pt; border-bottom: 2px solid #000; margin-top: 0.3in; }}
h3 {{ font-size: 12pt; margin-top: 0.2in; }}
ul {{ margin: 0.1in 0; }}
li {{ margin-bottom: 4pt; }}
</style></head><body>
{''.join(f'<p>{p}</p>' if not p.startswith('•') else f'<ul><li>{p[2:]}</li></ul>' for p in docx_text)}
</body></html>"""
        HTML(string=html_content).write_pdf(str(pdf_path))
        logger.info("Tailored resume PDF saved to %s", pdf_path)
        return pdf_path
    except Exception as e:
        logger.warning("PDF generation failed: %s — returning docx path", e)
        return docx_path
